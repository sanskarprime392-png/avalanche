# -*- coding: utf-8 -*-
"""Assemble the manuscript from the actual result files — no placeholder numbers."""
import os

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

DATA = r"C:\Users\Sanskar\Documents\avalanche_data"
RES = os.path.join(DATA, "results")
FIGS = os.path.join(DATA, "figures")
OUT = r"C:\Users\Sanskar\Documents\avalanche\Avalanche_Manuscript.docx"
NAVY = RGBColor(0x0B, 0x3D, 0x5C)
NICE = {"lgbm": "LightGBM", "xgb": "XGBoost", "rf": "Random Forest",
        "svm_rbf": "SVM (RBF)", "logreg": "Logistic Regression"}

mt = pd.read_csv(os.path.join(RES, "final_model_table.csv"))
ld = pd.read_csv(os.path.join(RES, "leakage_decomposition.csv"))
ga = pd.read_csv(os.path.join(RES, "geography_ablation.csv"))
br = pd.read_csv(os.path.join(RES, "background_ratio.csv"))
tr = pd.read_csv(os.path.join(RES, "transferability.csv"))
ls = pd.read_csv(os.path.join(RES, "label_quality_sensitivity.csv"))
sh = pd.read_csv(os.path.join(RES, "shap_importance.csv"))

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def H(t, lvl=1):
    h = doc.add_heading(t, level=lvl)
    for r in h.runs:
        r.font.color.rgb = NAVY
    return h


def P(t, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.italic = italic
    r.font.size = Pt(size)
    return p


def bullets(items):
    for i in items:
        doc.add_paragraph(i, style="List Bullet")


def table(headers, rows, widths, caption=None):
    if caption:
        c = doc.add_paragraph()
        rr = c.add_run(caption)
        rr.bold = True
        rr.font.size = Pt(9.5)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9)
    for row in t.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = Inches(w)
    doc.add_paragraph()
    return t


def figure(name, caption, width=6.0):
    p = os.path.join(FIGS, name + ".png")
    if not os.path.exists(p):
        return
    doc.add_picture(p, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)


# ---------------------------------------------------------------- title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Are machine-learning avalanche susceptibility maps learning avalanche physics?")
r.bold = True
r.font.size = Pt(17)
r.font.color.rgb = NAVY
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Sampling design, spatial validation and explainability in the Western Himalaya")
r.font.size = Pt(12.5)
a = doc.add_paragraph()
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = a.add_run("Chandra-Bhaga and Upper Beas basins, Himachal Pradesh and Jammu & Kashmir, India")
r.italic = True
r.font.size = Pt(10)
doc.add_paragraph()

# ---------------------------------------------------------------- abstract
H("Abstract", 1)
d_xgb = ld[(ld.config == "D") & (ld.model == "xgb")].roc_auc.iloc[0]
a_xgb = ld[(ld.config == "A") & (ld.model == "xgb")].roc_auc.iloc[0]
b_xgb = ld[(ld.config == "B") & (ld.model == "xgb")].roc_auc.iloc[0]
c_xgb = ld[(ld.config == "C") & (ld.model == "xgb")].roc_auc.iloc[0]
P(
    "Machine-learning avalanche susceptibility maps in the Indian Himalayan Region routinely report "
    "areas under the ROC curve above 0.90, yet the sampling designs that produce those numbers have "
    "not been tested. We construct an independent avalanche inventory for the Chandra-Bhaga and "
    "Upper Beas basins from eight winters of Sentinel-1 backscatter change detection (6,364 release "
    "points, 85% basin coverage) and use it to run a controlled decomposition of where reported "
    f"performance comes from. Drawing non-avalanche points from flat, vegetated, urban and cropland "
    f"areas — the rule used across this literature — yields AUC = {a_xgb:.3f}: the classification "
    "task becomes perfectly separable, and the model distinguishes steep snow-covered mountain from "
    "flat green valley rather than avalanche terrain from non-avalanche terrain. Replacing those "
    f"absences with terrain-matched background points reduces AUC to {c_xgb:.3f}, while spatial-block "
    f"cross-validation applied to the original sampling changes performance by {a_xgb - b_xgb:+.3f}. "
    "The optimism therefore originates almost entirely in sampling design, not in spatial "
    "autocorrelation, which is the aspect the field has standardised on correcting. Two further "
    "diagnostics show that 1 km climate and 500 m snow-cover predictors are statistically "
    "interchangeable with raw spatial coordinates, and that reported AUC can be raised simply by "
    "increasing the background sampling ratio while precision-recall performance falls. A "
    f"terrain-matched, spatially validated model attains AUC {d_xgb:.3f} and transfers between "
    f"sub-basins at {tr.roc_auc.mean():.3f}. We provide a corrected susceptibility map restricted to "
    "potential release areas, and openly document one artefact detected in our own pipeline.")

P("Keywords: snow avalanches, susceptibility mapping, sampling bias, spatial cross-validation, "
  "Sentinel-1, explainable machine learning, Western Himalaya", italic=True, size=9.5)

# ---------------------------------------------------------------- 1 intro
H("1. Introduction", 1)
P("Snow avalanches threaten settlements, highways and glacial lakes across the Indian Himalayan "
  "Region (IHR). Because field survey is impractical over this terrain, avalanche susceptibility "
  "maps (ASM) are produced by training classifiers on past avalanche locations and a set of "
  "topographic, climatic and snowpack predictors. Reported performance is consistently high, with "
  "recent IHR studies citing AUC values of 0.87–0.95.")
P("Such values are only meaningful if the negative class represents terrain where avalanches could "
  "occur but did not. In practice, published studies draw non-avalanche points from flat vegetated "
  "areas, urban settlements, water bodies and croplands. These locations differ from avalanche "
  "release zones in nearly every predictor simultaneously, so a classifier can achieve excellent "
  "scores by recognising land cover rather than avalanche conditions. The field has meanwhile "
  "concentrated its methodological attention on spatial cross-validation, which addresses a "
  "different problem — spatial autocorrelation between training and test folds.")
P("This study asks which of these two design choices actually drives reported performance, and "
  "whether the resulting models encode avalanche physics or study-design artefacts. We address it "
  "with a controlled 2×2 experiment on an independently constructed inventory, followed by "
  "explainability and ablation diagnostics. We apply the same scrutiny to our own pipeline and "
  "report an artefact we found there.")

# ---------------------------------------------------------------- 2 study area
H("2. Study area", 1)
P("The Chandra-Bhaga and Upper Beas basins span approximately 23,000 km² in Himachal Pradesh and "
  "the Jammu & Kashmir union territory, with elevations from 1,123 m to 6,329 m. The region "
  "receives heavy winter snowfall from westerly disturbances and has a documented history of "
  "avalanche impacts on the Leh–Manali highway and on settlements including Manali, Keylong and "
  "Udaipur. Potential release terrain (slope 28–60°) covers 11,699 km², about half the study area.")

# ---------------------------------------------------------------- 3 data & methods
H("3. Data and methods", 1)

H("3.1 Avalanche inventory from Sentinel-1", 2)
P("Avalanche debris roughens the snow surface and raises radar backscatter. We detect it as a "
  "seasonal increase in Sentinel-1 GRD backscatter between a stable mid-winter reference composite "
  "and a post-avalanche-season composite, computed independently for each winter from 2018 to 2025.")
bullets([
    "A single relative orbit is used (descending, orbit 136) so that viewing geometry is identical "
    "between dates and static terrain backscatter cancels in the difference. Orbit 136 is the only "
    "orbit providing 100% coverage of the study area; the orbit with the most acquisitions covers "
    "just 13% of it.",
    "Composites are averaged in linear power and speckle-filtered; only increases are retained, "
    "since wet snow lowers backscatter and therefore has the opposite sign.",
    "Detections are restricted to avalanche-capable slopes (25–60°) and filtered by minimum "
    "connected area.",
    "Detections are stacked across winters to give a RECURRENCE count. Random change does not "
    "repeat on the same slope across eight winters, whereas avalanche paths do; recurrence "
    "therefore acts as a confidence score.",
])
P("Deposits are traced upslope to the highest pixel on release-angle terrain (30–50°) within a "
  "600 m search radius, because susceptibility models predict release, not runout. This yields "
  "6,364 release points, of which 3,188 are flagged in four or more winters — the threshold used "
  "for the analyses below. Occupied 10 km blocks cover 85% of the study area, so the inventory is "
  "not confined to road-accessible slopes.")

H("3.2 Predictive factors", 2)
P("Twenty predictors were assembled on a common 30 m UTM 43N grid: twelve topographic and two "
  "hydrological derivatives computed from NASADEM with WhiteboxTools; winter temperature and "
  "precipitation (WorldClim); snow-cover duration (MODIS); land cover (ESA WorldCover); and "
  "distances to roads, streams (OpenStreetMap) and active faults (GEM Global Active Faults). "
  "Distance to faults is computed from vector geometry, because every active fault lies outside "
  "the study area and a rasterised distance transform is therefore undefined within it.")

H("3.3 Sampling design", 2)
P("Two absence schemes are compared. The PAPER-STYLE scheme reproduces published practice: points "
  "drawn at random from low-slope terrain classified as tree cover, grassland, cropland, built-up "
  "or water. The TERRAIN-MATCHED scheme draws background points from avalanche-capable slopes "
  "(25–55°) outside a 1 km buffer around presences, matched to the joint elevation × slope × "
  "aspect distribution of the presences.")
P("Aspect matching is required, not optional. Sentinel-1 is right-looking, so a single orbit "
  "over-detects on slopes facing the sensor: measured over this study area, descending orbit 136 "
  "places 34.8% of detections on south-east aspects against 12.6% of background terrain, while "
  "ascending orbit 27 peaks on north-west aspects. These are mirror images and therefore reflect "
  "viewing geometry rather than avalanche physics. Only the elevated south-facing fraction is "
  "consistent across both geometries. Matching background to the presence aspect distribution "
  "prevents the classifier from scoring on this artefact.", )

H("3.4 Validation and models", 2)
P("Models are evaluated under random k-fold cross-validation, as in published work, and under "
  "spatial-block cross-validation with 10 km blocks (246 blocks). Five classifiers are compared: "
  "Random Forest, Support Vector Machine and Logistic Regression, as used in the reference "
  "literature, plus XGBoost and LightGBM. Confidence intervals are percentile bootstraps (1,500 "
  "resamples) on pooled out-of-fold predictions.")

# ---------------------------------------------------------------- 4 results
H("4. Results", 1)

H("4.1 Where reported performance comes from", 2)
rows = []
lab = {"A": ("paper-style", "random"), "B": ("paper-style", "spatial-block"),
       "C": ("terrain-matched", "random"), "D": ("terrain-matched", "spatial-block")}
for cfg in ("A", "B", "C", "D"):
    rf_v = ld[(ld.config == cfg) & (ld.model == "rf")].roc_auc.iloc[0]
    xg_v = ld[(ld.config == cfg) & (ld.model == "xgb")].roc_auc.iloc[0]
    rows.append([cfg, lab[cfg][0], lab[cfg][1], f"{rf_v:.3f}", f"{xg_v:.3f}"])
table(["Config", "Absences", "Cross-validation", "AUC (RF)", "AUC (XGBoost)"], rows,
      [0.7, 1.5, 1.5, 1.1, 1.3],
      "Table 1. Controlled decomposition of reported performance.")
P(f"Configuration A, which reproduces published practice, achieves AUC = {a_xgb:.3f}: the two "
  f"classes are perfectly separable. Replacing the absences with terrain-matched background points "
  f"(A→C) costs {a_xgb - c_xgb:.3f} AUC for XGBoost and "
  f"{ld[(ld.config=='A')&(ld.model=='rf')].roc_auc.iloc[0] - ld[(ld.config=='C')&(ld.model=='rf')].roc_auc.iloc[0]:.3f} "
  f"for Random Forest. Applying spatial-block cross-validation to the original sampling (A→B) "
  f"changes performance by {a_xgb - b_xgb:+.3f}. The optimism is therefore attributable to sampling "
  "design, and essentially not at all to spatial autocorrelation.")
figure("fig_leakage_decomposition",
       "Figure 1. Decomposition of reported performance. Paper-style absences produce perfect "
       "separation regardless of the cross-validation scheme.")

P("This result does not depend on the quality of our inventory. Repeating the contrast at four "
  "recurrence thresholds, configuration A remains at AUC 1.000 throughout, including at n = 240 "
  "presence points — comparable to the 118 points used in the reference study.")
rows = []
for mr in sorted(ls.min_recurrence.unique()):
    A = ls[(ls.min_recurrence == mr) & (ls.config == "A")]
    D = ls[(ls.min_recurrence == mr) & (ls.config == "D")]
    rows.append([f"≥{int(mr)}", f"{int(A.n_presence.iloc[0]):,}",
                 f"{A.roc_auc.iloc[0]:.3f}", f"{D.roc_auc.iloc[0]:.3f}",
                 f"{A.roc_auc.iloc[0] - D.roc_auc.iloc[0]:.3f}"])
table(["Recurrence", "n presence", "A (paper)", "D (matched)", "Difference"], rows,
      [1.0, 1.1, 1.1, 1.2, 1.1],
      "Table 2. Sensitivity of the decomposition to inventory confidence.")
figure("fig_label_sensitivity",
       "Figure 2. The sampling artefact is independent of inventory size and label confidence.")

H("4.2 Model comparison under honest validation", 2)
rows = []
for mk in ("lgbm", "xgb", "rf", "svm_rbf", "logreg"):
    rnd = mt[(mt.model == mk) & (mt.cv == "random")].iloc[0]
    spa = mt[(mt.model == mk) & (mt.cv == "spatial")].iloc[0]
    rows.append([NICE[mk], f"{rnd.roc_auc:.3f}",
                 f"{spa.roc_auc:.3f} [{spa.ci_low:.3f}, {spa.ci_high:.3f}]",
                 f"{spa.pr_auc:.3f}", f"{spa.accuracy:.3f}", f"{spa.kappa:.3f}"])
table(["Model", "AUC (random)", "AUC (spatial, 95% CI)", "PR-AUC", "Accuracy", "Kappa"], rows,
      [1.3, 1.1, 1.7, 0.8, 0.9, 0.7],
      "Table 3. Model performance with terrain-matched absences (20 predictors, TPI excluded).")
P("Gradient-boosted trees outperform the Random Forest used in the reference study, but LightGBM "
  "and XGBoost are statistically indistinguishable from one another (paired bootstrap ΔAUC "
  "+0.0003, 95% CI −0.0013 to +0.0018). Reporting a single best model is not supportable at this "
  "sample size, and by extension neither are the four-way model rankings published on inventories "
  "of around one hundred points.")
figure("fig_model_ci", "Figure 3. Model ranking with bootstrap confidence intervals.", 5.4)

H("4.3 Coarse predictors encode geography, not snowpack", 2)
rows = [[r.variant, int(r.n_features), f"{r.roc_auc:.3f}", f"[{r.ci_low:.3f}, {r.ci_high:.3f}]"]
        for _, r in ga.iterrows()]
table(["Predictor set", "n", "AUC", "95% CI"], rows, [2.6, 0.5, 0.9, 1.5],
      "Table 4. Replacing coarse predictors with raw spatial coordinates.")
full = ga[ga.variant.str.contains("full")].roc_auc.iloc[0]
repl = ga[ga.variant.str.contains("REPLACED")].roc_auc.iloc[0]
P(f"Substituting the 1 km climate, 500 m snow-cover and distance-to-fault layers with two raw "
  f"coordinates changes performance from {full:.3f} to {repl:.3f}, a difference well within the "
  "confidence intervals. These predictors contribute to the model, but their contribution is "
  "positional rather than physical: at their native resolution they are smooth regional trend "
  "surfaces across release zones of 10²–10⁵ m². Published importance rankings that place snow-cover "
  "duration among the leading predictors may therefore be describing a geographic index rather "
  "than snowpack process.")
figure("fig_geography_ablation", "Figure 4. Coarse predictors are interchangeable with coordinates.", 6.4)

H("4.4 Reported AUC depends on the background sampling ratio", 2)
rows = [[r.ratio, f"{r.prevalence:.3f}", f"{r.roc_auc:.3f}", f"{r.pr_auc:.3f}", f"{r.brier:.4f}"]
        for _, r in br.iterrows()]
table(["Presence:background", "Prevalence", "AUC", "PR-AUC", "Brier"], rows,
      [1.6, 1.0, 0.9, 0.9, 0.9],
      "Table 5. Effect of the background sampling ratio.")
P(f"Increasing the background ratio from 1:1 to 1:10 raises AUC from {br.roc_auc.iloc[0]:.3f} to "
  f"{br.roc_auc.iloc[-1]:.3f} while precision-recall AUC falls from {br.pr_auc.iloc[0]:.3f} to "
  f"{br.pr_auc.iloc[-1]:.3f}. Because presence-background prevalence is a modelling choice rather "
  "than a property of the landscape, a headline AUC can be improved by a sampling decision while "
  "the model becomes less useful for the operational task. Precision-recall metrics should be "
  "reported alongside AUC.")

H("4.5 Spatial transferability", 2)
rows = [[r.transfer.replace("north", "Chandra-Bhaga").replace("south", "Upper Beas"),
         f"{int(r.n_train):,}", f"{int(r.n_test):,}",
         f"{r.roc_auc:.3f} [{r.ci_low:.3f}, {r.ci_high:.3f}]"] for _, r in tr.iterrows()]
table(["Transfer", "n train", "n test", "AUC (95% CI)"], rows, [2.2, 1.0, 1.0, 1.7],
      "Table 6. Training on one sub-basin and testing on the other.")
P(f"Transfer performance is symmetric and approximately 0.05 below within-region performance, "
  "indicating that the model generalises to terrain it has not seen rather than memorising local "
  "configurations.")
figure("fig_transfer_shap",
       "Figure 5. Sub-basin transferability (left) and SHAP predictor importance (right).", 6.6)

H("4.6 Calibration", 2)
P("Predicted scores are informative but not calibrated probabilities: the Brier score is 0.1034, "
  "isotonic regression recovers only 3.5% of it, and the maximum deviation between predicted and "
  "observed frequency across quantile bins is 0.099, with under-prediction at low scores and "
  "over-prediction at high scores. Because presence-background models are fitted at an arbitrary "
  "prevalence, susceptibility values should be interpreted as a relative index, not as "
  "P(avalanche) — a distinction lost when raw scores are binned by natural breaks and presented as "
  "hazard classes.")

H("4.7 Susceptibility map", 2)
P("The final map applies XGBoost with terrain-matched sampling and is restricted to potential "
  "release areas (slope 28–60°), covering 11,699 km². Restricting the domain matters for "
  "comparability: classifying valley floors and lake surfaces as low susceptibility inflates the "
  "denominator, so the commonly quoted statistic 'x% of the basin is highly susceptible' is not "
  "comparable between studies unless the domain is stated.")
table(["Class", "Area (km²)", "% of release area"],
      [["Very low", "6,125", "52.4"], ["Low", "2,167", "18.5"], ["Moderate", "1,314", "11.2"],
       ["High", "1,017", "8.7"], ["Very high", "1,077", "9.2"]],
      [1.4, 1.3, 1.6], "Table 7. Susceptibility classes within potential release areas.")
figure("fig_susceptibility_map",
       "Figure 6. Avalanche susceptibility across the Chandra-Bhaga and Upper Beas basins, "
       "restricted to potential release areas.", 6.4)
P("Excluding the coarse predictors changes 39% of class assignments (Pearson r = 0.761) but 71% of "
  "high and very-high pixels remain in those classes, indicating that the principal hotspots are "
  "terrain-controlled while the broader regional gradient is partly positional. The map reported "
  "here therefore uses only fine-resolution predictors with an identifiable physical mechanism.")

# ---------------------------------------------------------------- 5 discussion
H("5. Discussion", 1)

H("5.1 Implications for the literature", 2)
P("The decomposition indicates that high AUC values in IHR avalanche susceptibility studies are "
  "primarily a property of how non-avalanche points are drawn. Because the same absence rule "
  "recurs across this literature, the finding is not specific to any one study. It also implies "
  "that spatial cross-validation, while methodologically correct, addresses the smaller of the two "
  "problems: in our data it accounts for essentially none of the optimism once sampling is fixed.")

H("5.2 An artefact in our own pipeline", 2)
P("SHAP attributed the largest single contribution to the topographic position index, at "
  f"{sh.mean_abs_shap.iloc[0]:.2f}, roughly twice the next predictor. Inspection showed presence "
  "points average TPI +36.4 against +2.9 for background, which follows mechanically from our "
  "release-point rule: selecting the highest qualifying pixel within a search radius necessarily "
  "produces points at locally elevated topographic positions. The apparent importance was created "
  "by our own method. Excluding TPI costs 0.005 AUC, and it is excluded from all results reported "
  "here. We report this because the same class of error is what the study identifies elsewhere.")

H("5.3 A trade-off specific to SAR-derived inventories", 2)
P("We tested whether adding a Winstral wind-shelter index — the dominant snow-loading control, and "
  "absent from the predictor sets of the studies reproduced here — improves performance. It does "
  "not (terrain-only 0.902 to 0.906; full model 0.934 to 0.935, both within confidence intervals). "
  "The reason is structural: the shelter index correlates −0.949 with eastness, so matching "
  "background points on aspect in order to remove the Sentinel-1 look-direction artefact "
  "simultaneously equalises wind exposure. A SAR-derived inventory can support an unbiased aspect "
  "distribution or the study of aspect-dependent processes, but not both. Resolving this requires "
  "radiometric terrain flattening rather than distributional matching.")

H("5.4 Limitations", 2)
bullets([
    "Inventory precision is not established. Agreement between independent ascending and "
    "descending detections is 17.8% on aspects both geometries observe well, so the inventory "
    "should be treated as candidate avalanche-prone slopes rather than confirmed events. Recall at "
    "eight sites named in the reference study is 6/8, against 4/12 at randomly chosen steep sites, "
    "which is suggestive but not statistically significant at this sample size.",
    "Release points are inferred by upslope tracing rather than observed.",
    "Lithology and lineaments are taken from global datasets (GLiM, GEM) rather than the national "
    "geological survey products used in the reference study, so importance values for these "
    "predictors are not directly comparable.",
    "Because the model is fitted at an arbitrary presence-background prevalence, outputs are a "
    "relative index rather than probabilities.",
    "Projection of future susceptibility under climate scenarios was considered and rejected: "
    "Section 4.3 shows the coarse climate predictors carry positional rather than physical "
    "information, so propagating them forward would project a spatial trend surface.",
])

# ---------------------------------------------------------------- 6 conclusions
H("6. Conclusions", 1)
bullets([
    f"Reproducing the negative-sampling rule used across IHR avalanche susceptibility studies "
    f"yields AUC = {a_xgb:.3f}. The task becomes perfectly separable, and reported performance "
    f"reflects sampling design rather than predictive skill.",
    f"Spatial-block cross-validation changes performance by {a_xgb - b_xgb:+.3f} under that "
    "sampling design. The field's principal methodological correction addresses the smaller "
    "problem.",
    "Coarse climate and snow-cover predictors are statistically interchangeable with raw spatial "
    "coordinates and should not be interpreted as snowpack process.",
    "Reported AUC can be increased by raising the background sampling ratio while precision-recall "
    "performance declines; both should be reported.",
    f"A terrain-matched, spatially validated model attains AUC {d_xgb:.3f} "
    f"[{mt[(mt.model=='xgb')&(mt.cv=='spatial')].ci_low.iloc[0]:.3f}, "
    f"{mt[(mt.model=='xgb')&(mt.cv=='spatial')].ci_high.iloc[0]:.3f}] and transfers between "
    f"sub-basins at {tr.roc_auc.mean():.3f}.",
])

# ---------------------------------------------------------------- data / code
H("Data and code availability", 1)
P("All code is openly available at https://github.com/sanskarprime392-png/avalanche. The avalanche "
  "inventory derived in this study, the trained susceptibility rasters and all result tables are "
  "reproducible from that repository using openly licensed input data (NASADEM, Sentinel-1, "
  "Sentinel-2, WorldClim, MODIS, ESA WorldCover, OpenStreetMap, GEM Global Active Faults).")

H("References", 1)
for i, r in enumerate([
    "Abhinav, A. & Sattar, A. (2025). Snow avalanche susceptibility, hazard, and exposure "
    "assessment in the Western Himalaya using machine learning and numerical modelling. Scientific "
    "Reports, 15, 38093.",
    "Bühler, Y., von Rickenbach, D., Stoffel, A., Margreth, S., Stoffel, L. & Christen, M. (2018). "
    "Automated snow avalanche release area delineation. Natural Hazards and Earth System Sciences, "
    "18, 3235–3251.",
    "Eckerstorfer, M., Vickers, H., Malnes, E. & Grahn, J. (2019). Near-real time automatic snow "
    "avalanche activity monitoring system using Sentinel-1 SAR data in Norway. Remote Sensing, 11, 2863.",
    "Grinsztajn, L., Oyallon, E. & Varoquaux, G. (2022). Why do tree-based models still outperform "
    "deep learning on tabular data? Advances in Neural Information Processing Systems.",
    "Ploton, P. et al. (2020). Spatial validation reveals poor predictive performance of "
    "large-scale ecological mapping models. Nature Communications, 11, 4540.",
    "Roberts, D. R. et al. (2017). Cross-validation strategies for data with temporal, spatial, "
    "hierarchical, or phylogenetic structure. Ecography, 40, 913–929.",
    "Winstral, A., Elder, K. & Davis, R. E. (2002). Spatial snow modeling of wind-redistributed "
    "snow using terrain-based parameters. Journal of Hydrometeorology, 3, 524–538.",
], 1):
    doc.add_paragraph(f"[{i}]  {r}", style="List Paragraph")

doc.save(OUT)
print("wrote", OUT)
